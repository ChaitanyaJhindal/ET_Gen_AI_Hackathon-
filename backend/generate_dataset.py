import os
import pandas as pd
from datetime import datetime, timedelta
import random

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
except ImportError:
    print("Missing libraries. Run: pip install pandas python-docx reportlab openpyxl")
    exit(1)

out_dir = "../Shakti_Hackathon_Dataset"
os.makedirs(out_dir, exist_ok=True)

print(f"Generating massive dataset in {out_dir}...")

# 1. Generate PDFs (Equipment Manuals)
def create_pdf_manual(filename, title, content_sections):
    filepath = os.path.join(out_dir, filename)
    doc = SimpleDocTemplate(filepath, pagesize=letter)
    styles = getSampleStyleSheet()
    h1 = styles['Heading1']
    h2 = styles['Heading2']
    normal = styles['Normal']
    
    story = []
    story.append(Paragraph(title, h1))
    story.append(Spacer(1, 0.2 * inch))
    
    for section_title, text_blocks in content_sections.items():
        story.append(Paragraph(section_title, h2))
        story.append(Spacer(1, 0.1 * inch))
        if isinstance(text_blocks, list):
            for block in text_blocks:
                story.append(Paragraph(block, normal))
                story.append(Spacer(1, 0.1 * inch))
        else:
            story.append(Paragraph(text_blocks, normal))
            story.append(Spacer(1, 0.1 * inch))
            
    doc.build(story)
    print(f"Created {filename}")

# P-101
create_pdf_manual(
    "P-101_Feedwater_Pump_Manual.pdf",
    "P-101 High-Pressure Feedwater Pump - Service & Operations Manual",
    {
        "1. Overview": "The P-101 Feedwater Pump is a multi-stage centrifugal pump designed for high-pressure boiler feed applications. It supplies continuous water flow to the primary steam generation unit.",
        "2. Specifications": [
            "Max Flow Rate: 1500 GPM",
            "Operating Pressure: 2500 PSI",
            "RPM: 3550",
            "Bearing Type: Thrust and Journal (Hydrodynamic)",
            "Lubrication: Forced oil circulating system (ISO VG 46)"
        ],
        "3. Operating Thresholds": "Vibration must not exceed 0.25 in/s RMS. Oil temperature must remain between 40C and 65C. If oil temperature exceeds 70C, immediate shutdown is required.",
        "4. Troubleshooting Guide": [
            "Symptom: High Vibration. Cause: Bearing wear, misalignment, or cavitation. Action: Check suction pressure, inspect bearings, perform laser alignment.",
            "Symptom: Low Flow. Cause: Impeller wear or suction blockage. Action: Clean suction strainer, inspect impeller.",
            "Symptom: High Bearing Temperature. Cause: Inadequate lubrication. Action: Check lube oil pump, clean oil filters."
        ]
    }
)

# C-305
create_pdf_manual(
    "C-305_Compressor_Service_Manual.pdf",
    "C-305 Centrifugal Gas Compressor - Technical Manual",
    {
        "1. Equipment Summary": "C-305 is a 4-stage centrifugal gas compressor responsible for compressing raw syngas for the downstream synthesis loop.",
        "2. Key Parameters": [
            "Inlet Pressure: 45 PSI",
            "Discharge Pressure: 1200 PSI",
            "Surge Margin: Minimum 10%",
            "Drive: 5000 HP Synchronous Electric Motor"
        ],
        "3. Surge Prevention": "Surge is a catastrophic aerodynamic instability. The anti-surge valve (ASV) must open automatically if the operating point crosses the surge control line. Flow anomalies from upstream components (like Heat Exchangers) can trigger premature surge.",
        "4. Maintenance Intervals": "Vibration probes (Eddy Current) calibrated annually. Dry gas seals inspected every 24 months. Lube oil analyzed monthly."
    }
)

# HX-401
create_pdf_manual(
    "HX-401_Heat_Exchanger_Specs.pdf",
    "HX-401 Shell & Tube Heat Exchanger - Design & Maintenance Specs",
    {
        "1. Design Specifications": "HX-401 is a TEMA Type AES shell and tube heat exchanger. It cools hot syngas before it enters compressor C-305.",
        "2. Fouling Mechanics": "Mineral deposits from untreated cooling water can foul the tube side, leading to reduced heat transfer efficiency and an increased pressure drop.",
        "3. Operational Impacts": "If the pressure drop across HX-401 exceeds 15 PSI, the downstream compressor (C-305) will experience reduced suction pressure, pushing it closer to its surge limit.",
        "4. Cleaning Procedures": "Perform hydro-blasting or chemical cleaning if the differential pressure exceeds 15 PSI or if the approach temperature exceeds 20 degrees F."
    }
)

# 2. Generate Excel (Massive Maintenance Log)
def create_excel_logs():
    dates = [datetime(2022, 1, 1) + timedelta(days=i) for i in range(1000)]
    
    assets = ["P-101", "C-305", "HX-401", "G-502", "V-102", "M-201"]
    types = ["Preventive", "Corrective", "Inspection", "Calibration"]
    
    data = []
    for _ in range(500):
        asset = random.choice(assets)
        date = random.choice(dates)
        log_type = random.choice(types)
        
        desc = ""
        if asset == "HX-401" and log_type == "Corrective":
            desc = "Cleaned heavily fouled tubes. Pressure drop was critical at 18 PSI."
        elif asset == "C-305" and log_type == "Inspection":
            desc = "Anti-surge valve stroke test passed. Checked dry gas seal leakage."
        elif asset == "P-101" and log_type == "Corrective":
            desc = "Replaced degraded thrust bearing. Oil sample showed high metallic particulates."
        else:
            desc = f"Routine {log_type.lower()} maintenance performed as per schedule."
            
        data.append({
            "Date": date.strftime("%Y-%m-%d"),
            "Asset_ID": asset,
            "Maintenance_Type": log_type,
            "Description": desc,
            "Technician": f"Tech-{random.randint(10, 99)}",
            "Downtime_Hours": random.randint(1, 24)
        })
        
    data.append({
        "Date": "2024-03-15",
        "Asset_ID": "HX-401",
        "Maintenance_Type": "Corrective",
        "Description": "Emergency cleaning of HX-401. Severe mineral fouling caused a 22 PSI pressure drop across the exchanger, starving the downstream compressor.",
        "Technician": "Tech-Lead",
        "Downtime_Hours": 48
    })
    data.append({
        "Date": "2024-03-16",
        "Asset_ID": "C-305",
        "Maintenance_Type": "Corrective",
        "Description": "Compressor surged repeatedly due to low suction pressure caused by HX-401 blockage. Replaced damaged labyrinth seals.",
        "Technician": "Tech-Lead",
        "Downtime_Hours": 72
    })
    
    df = pd.DataFrame(data)
    df = df.sort_values(by="Date")
    filepath = os.path.join(out_dir, "Plant_Maintenance_Log_2022_2024.xlsx")
    df.to_excel(filepath, index=False)
    print(f"Created Excel Log")

create_excel_logs()

# 3. Generate Word Docs (Narrative Reports)
def create_word_doc(filename, title, paragraphs):
    filepath = os.path.join(out_dir, filename)
    doc = Document()
    doc.add_heading(title, 0)
    
    for p in paragraphs:
        doc.add_paragraph(p)
        
    doc.save(filepath)
    print(f"Created {filename}")

create_word_doc(
    "Incident_Report_C-305_Surge.docx",
    "Incident Report: Cascading Surge Failure on Compressor C-305",
    [
        "Date of Incident: March 16, 2024",
        "Author: Lead Reliability Engineer",
        "Executive Summary:",
        "On March 16, the main centrifugal compressor C-305 experienced three consecutive surge events resulting in an emergency trip and severe damage to the internal labyrinth seals.",
        "Root Cause Analysis:",
        "The root cause of the compressor surge was not an internal compressor fault, but rather a cascading failure originating upstream at the shell-and-tube heat exchanger HX-401.",
        "Over the previous 6 months, poor water treatment led to severe mineral fouling inside the tubes of HX-401. This fouling restricted gas flow, causing the differential pressure across the heat exchanger to spike to 22 PSI (the operational threshold is 15 PSI).",
        "Because HX-401 feeds directly into C-305, this severe pressure drop starved the compressor of suction pressure. When the suction pressure dropped below the critical threshold, the operating point of C-305 shifted drastically to the left on the performance map, crossing the surge line.",
        "Although the anti-surge valve (ASV) attempted to open, the pressure drop was too sudden. The compressor surged, causing intense aerodynamic buffeting that destroyed the labyrinth seals.",
        "Remediation:",
        "HX-401 tubes were hydro-blasted. C-305 seals were replaced. To prevent recurrence, a differential pressure alarm has been added to HX-401 to warn operators long before the compressor is starved."
    ]
)

create_word_doc(
    "Vibration_Analysis_P-101.docx",
    "Vibration Analysis Report: Feedwater Pump P-101",
    [
        "Date of Analysis: April 10, 2024",
        "Analyst: Vibration Specialist Team",
        "Findings:",
        "Routine condition monitoring on P-101 revealed an alarming upward trend in overall vibration levels. The RMS velocity on the outboard bearing reached 0.38 in/s, significantly exceeding the 0.25 in/s threshold defined in the service manual.",
        "Spectral analysis shows a dominant peak at 1X running speed, with notable harmonics at 2X and 3X. This signature strongly suggests severe shaft misalignment.",
        "Further investigation of the lube oil reservoir revealed metallic flakes, confirming that the sustained vibration from the misalignment has initiated hydrodynamic bearing degradation.",
        "Recommendation:",
        "P-101 must be taken offline immediately. The bearings require full replacement, and a precision laser alignment must be conducted before returning the pump to service to prevent catastrophic failure."
    ]
)

# 4. Generate Text Files (Shift Logs)
def create_text_file(filename, text):
    filepath = os.path.join(out_dir, filename)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f"Created {filename}")

create_text_file(
    "Shift_Supervisor_Log_Week11.txt",
    """SHIFT SUPERVISOR LOG - WEEK 11, 2024

Monday: Normal operations.
Tuesday: Noticed HX-401 differential pressure creeping up. Logged at 16 PSI. Will monitor.
Wednesday: HX-401 DP now at 19 PSI. Operator noted C-305 compressor sounding slightly rough.
Thursday: C-305 tripped! Massive surge event at 14:00 hours. Maintenance team deployed. HX-401 DP was pegged at 22 PSI right before the trip. Looks like the heat exchanger choked the compressor.
Friday: Plant offline for C-305 seal replacement and HX-401 cleaning.
"""
)

print("Massive dataset generation complete!")
